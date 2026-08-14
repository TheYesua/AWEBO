"""Normalización de los códigos de currículo que devuelve el modelo.

POR QUÉ ESTO NO SE ARREGLÓ EN EL PROMPT
----------------------------------------
Se intentó, el 14/08/2026. El ejemplo del prompt de conexión curricular llevaba
`"CE1"` —la convención de la literatura didáctica, no la del boletín— y se
sustituyó por un marcador sin valor concreto.

El resultado, comprobado generando la misma SdA dos veces: **el fallo cambió de
idioma**. Antes del arreglo, la versión castellana ponía el prefijo y la
catalana no; después, al revés. Un arreglo que mueve el fallo de sitio en vez
de quitarlo demuestra que la causa no era el ejemplo, o no solo: es
variabilidad del modelo.

Pedirle a un modelo que respete un formato es razonable. Confiar en que lo haga
**siempre**, cuando de ello depende que un criterio de evaluación se enlace o
se tire, no lo es. Por eso la normalización vive aquí, que es determinista.
"""
from __future__ import annotations

from pathlib import Path

import pytest

from app.services.enlaces_curriculares import variantes_de_codigo


class TestElCodigoTalCualVaPrimero:
    """El orden no es un detalle: decide qué gana si el catálogo es ambiguo."""

    def test_el_original_siempre_encabeza(self):
        """Si algún día un catálogo usa de verdad el prefijo, el literal tiene
        que ganar. Normalizar primero y preguntar después convertiría un código
        correcto en otro que también existe, y el resultado parecería bien."""
        assert variantes_de_codigo("CE1")[0] == "CE1"
        assert variantes_de_codigo("1")[0] == "1"

    def test_una_cadena_vacia_no_produce_variantes(self):
        assert variantes_de_codigo("") == []
        assert variantes_de_codigo("   ") == []


class TestElPrefijoQueSeInventaElModelo:
    @pytest.mark.parametrize("escrito, esperado", [
        ("CE1", "1"),
        ("CE 3", "3"),
        ("CE-3", "3"),
        ("C.E. 5", "5"),
        ("ce7", "7"),
    ])
    def test_se_prueba_tambien_sin_prefijo(self, escrito, esperado):
        """Ni el BOE ni el DOGC numeran las competencias con prefijo: son «1»,
        «2»… El modelo escribe «CE1» porque es lo habitual en la literatura
        didáctica, y sin esto la competencia se descartaba en silencio."""
        assert esperado in variantes_de_codigo(escrito)

    def test_y_al_reves_por_si_el_catalogo_lo_llevara(self):
        """La simetría no es adorno: hoy ningún catálogo usa prefijo, pero el
        día que uno lo use, el modelo escribirá «1» y habrá que encontrarlo."""
        assert "CE1" in variantes_de_codigo("1")

    def test_un_codigo_de_saber_no_se_toca(self):
        """«A.3» no lleva prefijo de competencia. Añadirle uno buscaría una
        fila que no existe, y peor: podría existir en otra materia."""
        assert variantes_de_codigo("A.3") == ["A.3"]


class TestEspaciosDeMas:
    def test_se_prueba_tambien_compacto(self):
        """«1. 1» por «1.1». El modelo parte el código al ajustar el texto."""
        assert "1.1" in variantes_de_codigo("1. 1")


class TestNoSeInventanCoincidencias:
    @pytest.mark.parametrize("codigo", ["CE", "criterio", "1.1.1.1", "—"])
    def test_lo_que_no_es_un_codigo_no_gana_variantes_raras(self, codigo):
        """Cada variante extra es una oportunidad de casar con la fila
        equivocada. Solo se generan cuando hay un patrón reconocible."""
        variantes = variantes_de_codigo(codigo)

        assert variantes[0] == codigo
        assert all(codigo in v or v in codigo or v.lstrip("CE") == codigo
                   for v in variantes), variantes


class TestUnaConexionIncompletaSeDice:
    """Que el documento exportado no disimule una sección que falta.

    Cada bloque se pintaba solo si tenía datos, así que una generación
    incompleta producía un documento que **parecía completo**: sin criterios,
    el PDF pasa de competencias a saberes y nada indica que falte nada.

    Pasó el 14/08/2026 con la misma SdA en catalán y en castellano: la
    castellana traía sus cuatro criterios y la catalana ninguno. No lo filtraba
    la exportación —pinta lo que hay en el JSONB, sin cribar—: el modelo no los
    devolvió.

    Desde aquí no se puede impedir, pero sí dejar de disimularlo.
    """

    @staticmethod
    def _texto(datos):
        from docx import Document

        from app.services import exportacion_service as ex

        doc = Document()
        ex._docx_conexion_curricular(doc, datos)
        return "\n".join(p.text for p in doc.paragraphs)

    COMPLETA = {
        "competencias": [{"codigo": "1", "justificacion": "x"}],
        "criterios": [{"codigo": "1.1", "competencia": "1", "justificacion": "y"}],
        "saberes": [{"codigo": "20.1", "justificacion": "z"}],
    }

    def test_sin_criterios_lo_avisa_y_los_nombra(self):
        datos = {k: v for k, v in self.COMPLETA.items() if k != "criterios"}

        texto = self._texto(datos)

        assert "incompleta" in texto
        assert "criterios de evaluación" in texto

    def test_una_conexion_completa_no_lleva_aviso(self):
        """Un aviso que sale siempre deja de leerse."""
        assert "⚠" not in self._texto(self.COMPLETA)

    def test_si_faltan_dos_se_nombran_las_dos(self):
        texto = self._texto({"competencias": self.COMPLETA["competencias"]})

        assert "criterios de evaluación" in texto
        assert "saberes básicos" in texto

    def test_el_aviso_va_antes_del_contenido(self):
        """Al final del bloque se lo saltaría quien lea en diagonal."""
        datos = {k: v for k, v in self.COMPLETA.items() if k != "criterios"}

        texto = self._texto(datos)

        assert texto.index("incompleta") < texto.index("Competencias específicas")


class TestElJSONBSeQuedaConElCodigoDelBoletin:
    """Enlazar bien por dentro y enseñar mal por fuera no resuelve nada.

    `variantes_de_codigo` hace que «CE1» encuentre la fila «1», así que la SdA
    queda bien enlazada. Pero el documento exportado **no se pinta de los
    enlaces**: se pinta del JSONB, donde seguía lo que escribió el modelo. El
    docente veía «CE1» en su PDF y ese código no está en el decreto.
    """

    @staticmethod
    def _contenido():
        return {"conexion_curricular": {
            "competencias": [{"codigo": "CE1", "justificacion": "x"}],
            "criterios": [{"codigo": "1.1", "competencia": "CE1", "justificacion": "y"}],
            "saberes": [{"codigo": "20.1", "justificacion": "z"}],
        }}

    def test_el_codigo_citado_se_cambia_por_el_del_catalogo(self):
        from app.services.enlaces_curriculares import _reescribir_codigos

        contenido = self._contenido()

        assert _reescribir_codigos(contenido, "competencias", {"CE1": "1"})
        assert contenido["conexion_curricular"]["competencias"][0]["codigo"] == "1"

    def test_tambien_la_referencia_del_criterio_a_su_competencia(self):
        """Un criterio apunta a su competencia por código. Corregir uno y no el
        otro deja el documento con un criterio que referencia algo que ya no
        aparece en la tabla de arriba."""
        from app.services.enlaces_curriculares import _reescribir_codigos

        contenido = self._contenido()
        _reescribir_codigos(contenido, "criterios", {"CE1": "1"})

        assert contenido["conexion_curricular"]["criterios"][0]["competencia"] == "1"

    def test_lo_que_no_se_resolvio_se_queda_como_estaba(self):
        """Un código huérfano hay que poder verlo: es la señal de que el modelo
        se lo inventó. Normalizarlo lo escondería."""
        from app.services.enlaces_curriculares import _reescribir_codigos

        contenido = self._contenido()
        _reescribir_codigos(contenido, "saberes", {"CE1": "1"})

        assert contenido["conexion_curricular"]["saberes"][0]["codigo"] == "20.1"

    def test_sin_nada_que_cambiar_no_dice_que_cambio(self):
        """El valor de retorno decide si se reasigna el JSONB. Un `True` de más
        escribe en la base de datos en cada sincronización, sin motivo."""
        from app.services.enlaces_curriculares import _reescribir_codigos

        assert not _reescribir_codigos(self._contenido(), "competencias", {})

    def test_un_jsonb_retorcido_no_revienta(self):
        """Lo escribe un modelo: ha llegado como lista de cadenas, como dict, y
        ausente. Reventar aquí tumbaría el final de una generación correcta."""
        from app.services.enlaces_curriculares import _reescribir_codigos

        for basura in ({}, {"conexion_curricular": "texto"},
                       {"conexion_curricular": {"competencias": "no es lista"}},
                       {"conexion_curricular": {"competencias": ["CE1"]}}):
            assert _reescribir_codigos(basura, "competencias", {"CE1": "1"}) is False


class TestLasDosRutasDeExportacionAvisanIgual:
    """Hay DOS caminos de exportación, y el aviso se puso en uno.

    `_docx_conexion_curricular` pinta el .docx; `render_conexion_curricular`,
    un macro de `templates/exportacion/pdf.html`, pinta el PDF. Son código
    distinto que produce lo mismo.

    El 14/08/2026 se añadió el aviso de sección incompleta **solo al docx**, y
    en el PDF exportado no salía. Se dio por hecho que el problema era que el
    contenedor no había recargado el código; se recreó, y seguía sin salir.
    La causa era esta: se tocó una de las dos rutas.

    Este test recorre las dos. Añadir una tercera sin pasar por aquí volvería
    a dejar un camino mudo.
    """

    COMPLETA = {
        "competencias": [{"codigo": "1", "justificacion": "x"}],
        "criterios": [{"codigo": "1.1", "competencia": "1", "justificacion": "y"}],
        "saberes": [{"codigo": "20.1", "justificacion": "z"}],
    }

    @staticmethod
    def _docx(datos):
        from docx import Document

        from app.services import exportacion_service as ex

        doc = Document()
        ex._docx_conexion_curricular(doc, datos)
        return "\n".join(p.text for p in doc.paragraphs)

    @staticmethod
    def _pdf(datos):
        """El macro suelto, con su dependencia.

        Se extrae del fichero real y no se copia: si alguien cambia la
        plantilla, esto tiene que enterarse. Renderizar la plantilla entera
        exigiría `sa`, `usuario` y media aplicación.
        """
        from jinja2 import DictLoader, Environment

        fuente = Path("app/templates/exportacion/pdf.html").read_text(encoding="utf-8")

        def macro(nombre):
            i = fuente.index("{%- macro " + nombre)
            j = fuente.index("{%- endmacro -%}", i) + len("{%- endmacro -%}")
            return fuente[i:j]

        env = Environment(
            loader=DictLoader({"m": macro("tabla_curr") + "\n"
                               + macro("render_conexion_curricular")}),
            extensions=["jinja2.ext.i18n"],
        )
        env.install_null_translations(newstyle=True)
        return env.get_template("m").module.render_conexion_curricular(datos)

    @pytest.fixture(params=["docx", "pdf"])
    def pintar(self, request):
        return self._docx if request.param == "docx" else self._pdf

    def test_sin_criterios_avisa_y_los_nombra(self, pintar):
        datos = {k: v for k, v in self.COMPLETA.items() if k != "criterios"}

        texto = pintar(datos)

        assert "incompleta" in texto
        assert "criterios de evaluación" in texto

    def test_una_conexion_completa_no_lleva_aviso(self, pintar):
        """Un aviso que sale siempre deja de leerse."""
        assert "incompleta" not in pintar(self.COMPLETA)

    def test_no_nombra_las_secciones_que_si_estan(self, pintar):
        datos = {k: v for k, v in self.COMPLETA.items() if k != "criterios"}

        aviso = pintar(datos).split("⚠")[1].split("\n")[0]

        assert "saberes" not in aviso and "competencias" not in aviso
